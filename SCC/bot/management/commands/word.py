from bot.models import *
from django.core.management.base import BaseCommand
import docx

class Command(BaseCommand):
  	# Используется как описание команды обычно
    help = 'Implemented to Django application telegram bot setup command'
    
    def handle(self, *args, **options):
        doc = docx.Document()
        kwargs = list(options.items())[7:]
        kwargs = dict(kwargs)
        applications = Application.objects.filter(IsFinished=True, **kwargs)
        
        # добавляем таблицу 3x3
        table = doc.add_table(rows = applications.count() + 1, cols = 4)
        # применяем стиль для таблицы
        table.style = 'Table Grid'
        table.cell(0, 0).paragraphs[0].add_run("Коллектив").bold = True
        table.cell(0, 1).paragraphs[0].add_run("ФИО").bold = True
        table.columns[1].width = 2500000
        table.cell(0, 2).paragraphs[0].add_run("Номер телефона").bold = True
        table.columns[2].width = 1500000
        table.cell(0, 3).paragraphs[0].add_run("Дата").bold = True
        table.columns[3].width = 600000
        # заполняем таблицу данными
        for i in range(applications.count()):
            row = i + 1
            app = applications[i]
            print(app)
            cell = table.cell(row, 0)
            cell.text = str(app.Team.Name)
            cell = table.cell(row, 1)
            cell.text = str(app.Name)
            cell = table.cell(row, 2)
            cell.text = str(app.Phone)
            cell = table.cell(row, 3)
            cell.text = str(app.Date.Date)
        name = 'dw'
        for i in kwargs:
            name += kwargs[i]
        doc.save(f"static/docx/{name}.docx")

    def add_arguments(self, parser):
        parser.add_argument(
        '-d', 
        '--Date__Date',
        help='Фильтрация по дате'
        )