from django.shortcuts import render
from django.views import View
from bot.forms import *
from bot.models import *
from django.http import JsonResponse
from datetime import datetime, tzinfo
from django.utils.timezone import utc
import json
import docx
from django.contrib.auth import authenticate
from django.contrib.auth import login
from django.shortcuts import redirect
from django.http import QueryDict
from django.http import FileResponse
# Create your views here.

class Main(View):
    def get(self, request):
        if request.user.is_authenticated:
            names = []
            for group in request.user.groups.all():
                names.append(str(group))
            if request.user.is_superuser:
                d_args = {}
                args = {}
                t_args = {}
            else:
                d_args = {"dateteam__Team__Name__in": names}
                args = {"Team__Name__in": names}
                t_args = {"Name__in": names}
            context = {"applications": Application.objects.filter(IsFinished=True, **args), "team_form": TeamForm(), "teams": Team.objects.filter(**t_args), "Dates": Date.objects.filter(**d_args).order_by('Date'), "direction_form": DirectionForm(), "directions": reversed(Direction.objects.all())}
            return render(template_name='index.html', context = context, request=request)
        else:
            return redirect("/login")            

class Login(View):
    def get(self, request):
        return render(template_name='login.html', request=request)

    def post(self, request):
        username = request.POST["username"]
        password = request.POST["password"]
        user = authenticate(username = username, password = password)
        if user is not None:
            login(request, user)
            return redirect("/")
        else:
            return render(template_name='login.html', context = {"error": "Неправильный логин или пароль"}, request=request)
class Applications(View):
    def get(self, request):
        if request.user.is_authenticated:
            names = []
            for group in request.user.groups.all():
                print(group.name)
                names.append(str(group))
            if request.user.is_superuser:
                args = {}
            else:
                args = {"Team__Name__in": names}
            return JsonResponse({"applications": list(Application.objects.filter(IsFinished=True, **args).values('id', 'Name', 'Phone', 'Direction__Name', 'Team__Name', 'Date__Date'))})
        else:
            return redirect("/login")

class Directions(View):
    def get(self, request):
        if request.user.is_authenticated:
            return JsonResponse({"directions": list(Direction.objects.all().values('id', 'Name'))})

        else:
            return redirect("/login")

    def post(self, request):
        print(request.body)
        dir = Direction.objects.create(Name = json.loads(request.body)['name'])
        dir.save()
        return JsonResponse({"id": dir.id, "Name": dir.Name})

    def delete(self, request):
        return JsonResponse({'deleted': Direction.objects.get(id=request.GET.get('id')).delete()})

class Teams(View):
    def get(self, request):
        return JsonResponse({"teams": list(Team.objects.all().values('id', "Direction", "Name", "Picture", "Decription", "Place", "Vk", "Prompt", "Contacts", "Time"))})    
    
    def post(self, request):
        form = TeamForm(request.POST, request.FILES)
        print(form.is_valid())
        if form.is_valid():
            team = form.save()
            team.save()
            
            for dt in request.POST.getlist("Date"):
                print(dt)
                new_date = DateTeam.objects.create(Team = team, Date = Date.objects.get_or_create(Date=dt)[0])
                new_date.save()
            
            return JsonResponse({"team": {'id': team.id, "Direction": team.Direction.Name, "Name": team.Name, "Picture": team.Picture.url, "Decription": team.Decription, "Place": team.Place, "Vk": team.Vk, "Prompt": team.Prompt, "Contacts": team.Contacts, "Time": team.Time}})
        for field, error in form.errors.items():
            print(field, ": ", error)
        context = {"applications": Application.objects.filter(IsFinished=True), "team_form": form, "teams": Team.objects.all(), "direction_form": DirectionForm(), "directions": reversed(Direction.objects.all())}
        return render(template_name='index.html', context = context, request=request)
        

    def delete(self, request):
        return JsonResponse({'deleted': Team.objects.get(id=request.GET.get('id')).delete()})
        

def update_team(request):
    if request.method == "POST":
        team = Team.objects.get(id = request.POST.get('id'))

        if request.POST.get('Name'):
            team.Name = request.POST.get('Name')
        if request.POST.get('Decription'): 
            team.Decription = request.POST.get('Decription')
        if request.POST.get('Place'):
            team.Place = request.POST.get('Place')
        if request.POST.get('Time'):
            team.Time = request.POST.get('Time')
        if request.POST.get('Direction'):
            team.Direction = Direction.objects.get(id = request.POST.get('Direction'))
        if request.POST.get('Vk'):
            team.Vk = request.POST.get('Vk')
        if request.POST.get('Contacts'):
            team.Contacts = request.POST.get('Contacts')
        if request.POST.get('Prompt'):
            team.Prompt = request.POST.get('Prompt')      
        team.save()
        for dt in DateTeam.objects.filter(Team=team):
            print(dt.Date.dateteam_set.all())
            date = dt.Date
            dt.delete()
            if date.dateteam_set.all().count() == 0:
                date.delete()
        for dt in request.POST.getlist("Date"):
            print(dt)
            new_date = DateTeam.objects.create(Team = team, Date = Date.objects.get_or_create(Date = dt)[0])
            new_date.save()

        return JsonResponse({'status': "OK"})
    

def download_file(request):
    if request.user.is_authenticated and request.method == 'POST':
        args = {}
        if len(request.POST.getlist("Date"))>0:
            args["Date__Date__in"] = request.POST.getlist("Date")
        if len(request.POST.getlist("Teams"))>0:
            args["Team__Name__in"] = request.POST.getlist("Teams")
        doc = docx.Document()
        applications = Application.objects.filter(IsFinished=True, **args).order_by('Team__Name', "Date__Date")
        if applications.count() == 0:
            return render(template_name='no_applications.html', request=request)
        # добавляем таблицу 3x3
        table = doc.add_table(rows = applications.count() + 1, cols = 4)
        # применяем стиль для таблицы
        table.style = 'Table Grid'
        table.cell(0, 0).paragraphs[0].add_run("Коллектив").bold = True
        table.cell(0, 1).paragraphs[0].add_run("ФИО").bold = True
        table.columns[1].width = 2500000
        table.cell(0, 2).paragraphs[0].add_run("Номер телефона").bold = True
        table.columns[2].width = 1500000
        table.cell(0, 3).paragraphs[0].add_run("Дата").bold = True
        table.columns[3].width = 600000
        # заполняем таблицу данными
        for i in range(applications.count()):
            row = i + 1
            app = applications[i]
            cell = table.cell(row, 0)
            cell.text = str(app.Team.Name)
            cell = table.cell(row, 1)
            cell.text = str(app.Name)
            cell = table.cell(row, 2)
            cell.text = str(app.Phone)
            cell = table.cell(row, 3)
            cell.text = str(app.Date.Date)
        name = 'dw'
        doc.save(f"static/docx/{name}.docx")
        return FileResponse(open(f"static/docx/{name}.docx", 'rb'))
    return redirect('/')
